from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import os
import requests
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for .docx files
from langchain.docstore.document import Document
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from io import BytesIO
from typing import List
from openai import OpenAI

# Initialize FastAPI
app = FastAPI()

# Initialize OpenAI client
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# Function to parse PDFs
def parse_pdf(file: BytesIO) -> str:
    pdf_document = fitz.open(stream=file, filetype="pdf")
    text = ""
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text

# Function to parse DOCX files
def parse_docx(file: BytesIO) -> str:
    doc = docx.Document(file)
    full_text = []

    # Parse paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Parse tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            full_text.append("\t".join(row_data))

    return "\n".join(full_text)

# Function to determine file type and extract content
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    if filename.endswith(".pdf"):
        return parse_pdf(BytesIO(file_bytes))
    elif filename.endswith(".docx"):
        return parse_docx(BytesIO(file_bytes))
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file format for file: {filename}")

# Function to convert text into documents
def text_to_docs(text: str, filename: str) -> List[Document]:
    if isinstance(text, str):
        text = [text]
    page_docs = [Document(page_content=page) for page in text]
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    doc_chunks = []
    for doc in page_docs:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=4000,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=0,
        )
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={"page": doc.metadata["page"], "chunk": i}
            )
            doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc.metadata["filename"] = filename  # Add filename to metadata
            doc_chunks.append(doc)
    return doc_chunks

# Function to create an index using FAISS
def docs_to_index(docs, openai_api_key):
    index = FAISS.from_documents(docs, OpenAIEmbeddings(openai_api_key=openai_api_key))
    return index

# Function to generate an index from URLs and process the query
def process_files_and_query(file_urls, query):
    documents = []
    for url in file_urls:
        try:
            url = url.replace("localhost", "rails")
            response = requests.get(url)
            response.raise_for_status()
            file_content = response.content
            filename = url.split("/")[-1]  # Extract filename from URL
            text = extract_text_from_file(file_content, filename)
            documents += text_to_docs(text, filename)
        except requests.RequestException as e:
            raise HTTPException(status_code=400, detail=f"Error downloading file from {url}: {str(e)}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing file from {url}: {str(e)}")

    # Create the FAISS index
    index = docs_to_index(documents, client.api_key)

    # Search the vector database
    search_results = index.similarity_search(query, k=3)

    pdf_extract = "\n".join([result.page_content for result in search_results])

    # Generate a response using OpenAI GPT-3.5 turbo with streaming
    prompt_template = """
    You are a helpful Assistant who answers users' questions based on multiple contexts given to you.

    Keep your answer short and to the point.

    The evidence is the context of the PDF extract with metadata.

    Carefully focus on the metadata, especially 'filename' and 'page' whenever answering.

    Make sure to add filename and page number at the end of the sentence you are citing.

    Reply "Not applicable" if the text is irrelevant.

    The PDF content is:
    {pdf_extract}
    """

    prompt = prompt_template.format(pdf_extract=pdf_extract)

    # Use the OpenAI client to generate a streaming response
    stream = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": query}
        ],
        stream=True,
    )
    print(stream)

    response = []
    result = ""
    for chunk in stream:
      text = chunk.choices[0].delta.content
      if text:
          response.append(text)
          result = "".join(response).strip()

    return result

# Endpoint to process the query and files
class QueryRequest(BaseModel):
    query: str
    file_urls: List[str]

@app.post("/process/")
async def process_document(request: QueryRequest):
    try:
        answer = process_files_and_query(request.file_urls, request.query)
        return {"answer": answer}
    except Exception as e:
        # Catch the exception and return it as part of the answer field
        return JSONResponse(content={"answer": f"Error processing request: {str(e)}"}, status_code=200)
